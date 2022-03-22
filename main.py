from tkinter import *
from tkinter import ttk
import tkinter.messagebox
from docx import Document
from docx.oxml.ns import qn  # 设置字体
from docx.shared import Pt  # 设置字体大小
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 设置对象居中、对齐等
import math
from ttkbootstrap import Style

LOG_LINE_NUM = 0
style = Style(theme='minty')

class MY_GUI():
    def __init__(self,init_window_name):
        self.init_window_name = init_window_name


    #设置窗口
    def set_init_window(self):
        self.init_window_name.title("悬挑板计算小工具_v1.0")           #窗口名
        #self.init_window_name.geometry('320x160+10+10')                         #290 160为窗口大小，+10 +10 定义窗口弹出时的默认展示位置
        self.init_window_name.geometry('1068x681+10+10')
        #self.init_window_name["bg"] = "pink"                                    #窗口背景色，其他背景色见：blog.csdn.net/chl0000/article/details/7657887
        #self.init_window_name.attributes("-alpha",0.9)                          #虚化，值越小虚化程度越高
        #标签
        self.inputlabel = Label(self.init_window_name, text="信息输入：", font=("宋体", 12, "bold"))
        self.inputlabel.grid(row=0, column=0, sticky='e')      
        self.namelabel = Label(self.init_window_name, text="构件名称：")
        self.namelabel.grid(row=1, column=1, sticky='e')
        self.hlabel = Label(self.init_window_name, text="板厚h（mm）：")
        self.hlabel.grid(row=2, column=1, sticky='e')
        self.llabel = Label(self.init_window_name, text="挑出长度l（mm）：")
        self.llabel.grid(row=3, column=1, sticky='e')
        self.blabel = Label(self.init_window_name, text="计算宽度b（mm）：")
        self.blabel.grid(row=4, column=1, sticky='e')
        self.dllabel = Label(self.init_window_name, text="恒载dl（kN/m\u00b2）：")
        self.dllabel.grid(row=5, column=1, sticky='e')
        self.lllabel = Label(self.init_window_name, text="活载ll（kN/m\u00b2）：")
        self.lllabel.grid(row=6, column=1, sticky='e')
        self.bllabel = Label(self.init_window_name, text="边缘翻边线荷载bl(kN/m)：")
        self.bllabel.grid(row=7, column=1, sticky='e')
        self.cglabel = Label(self.init_window_name, text="混凝土强度等级c_g：")
        self.cglabel.grid(row=8, column=1, sticky='e')
        self.fclabel = Label(self.init_window_name, text="fc（N/mm\u00b2）：")
        self.fclabel.grid(row=9, column=1, sticky='e')
        self.ftlabel = Label(self.init_window_name, text="ft（N/mm\u00b2）：")
        self.ftlabel.grid(row=10, column=1, sticky='e')
        self.sglabel = Label(self.init_window_name, text="钢筋强度等级s_g：")
        self.sglabel.grid(row=11, column=1, sticky='e')
        self.fylabel = Label(self.init_window_name, text="fy（N/mm\u00b2）：")
        self.fylabel.grid(row=12, column=1, sticky='e')       
        self.pminlabel = Label(self.init_window_name, text="最小配筋率：")
        self.pminlabel.grid(row=13, column=1, sticky='e')       
        self.aslabel = Label(self.init_window_name, text="纵筋合力点至近边距离（mm）：")
        self.aslabel.grid(row=14, column=1, sticky='e')    
        self.y0label = Label(self.init_window_name, text="重要性系数：")
        self.y0label.grid(row=15, column=1, sticky='e')    
        self.bblabel = Label(self.init_window_name, text="是否计算内部梁：")
        self.bblabel.grid(row=16, column=1, sticky='e')    
        self.Lblabel = Label(self.init_window_name, text="梁宽（mm）：")
        self.Lblabel.grid(row=17, column=1, sticky='e')    
        self.Lhlabel = Label(self.init_window_name, text="梁高（mm）：")
        self.Lhlabel.grid(row=18, column=1, sticky='e')    
        self.LLlabel = Label(self.init_window_name, text="梁长（mm）：")
        self.LLlabel.grid(row=19, column=1, sticky='e')    
        self.nameText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.nameText.grid(row=1, column=2)
        self.nameText.insert(0,'雨棚1')
        self.hText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.hText.grid(row=2, column=2)
        self.hText.insert(0,'120')
        self.lText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.lText.grid(row=3, column=2)
        self.lText.insert(0,'1200')
        self.bText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.bText.grid(row=4, column=2)
        self.bText.insert(0,'1000')
        self.dlText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.dlText.grid(row=5, column=2)
        self.dlText.insert(0,'1.5')
        self.llText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.llText.grid(row=6, column=2)
        self.llText.insert(0,'2')
        self.blText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.blText.grid(row=7, column=2)
        self.blText.insert(0,'1')
        self.cgText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.cgText.grid(row=8, column=2)
        self.cgText.insert(0,'C30')
        self.fcText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.fcText.grid(row=9, column=2)
        self.fcText.insert(0,'14.3')
        self.ftText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.ftText.grid(row=10, column=2)
        self.ftText.insert(0,'1.43')
        self.sgText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.sgText.grid(row=11, column=2)
        self.sgText.insert(0,'HRB400')
        self.fyText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.fyText.grid(row=12, column=2)
        self.fyText.insert(0,'360')
        self.pminText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.pminText.grid(row=13, column=2)
        self.pminText.insert(0,'0.002')
        self.asText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.asText.grid(row=14, column=2)
        self.asText.insert(0,'20')
        self.y0Text = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.y0Text.grid(row=15, column=2)
        self.y0Text.insert(0,'1.0')

        self.bbselection = ttk.Combobox (self.init_window_name,  width=17, textvariable='')
        self.bbselection["value"] = ("是", "否")
        self.bbselection.grid(row=16, column=2)
        self.bbselection.current(0)

        self.LbText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.LbText.grid(row=17, column=2)
        self.LbText.insert(0,'200')
        self.LhText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.LhText.grid(row=18, column=2)
        self.LhText.insert(0,'500')
        self.LLText = Entry(self.init_window_name, width=20,textvariable="")  #原始数据录入框
        self.LLText.grid(row=19, column=2)
        self.LLText.insert(0,'2400')
        
        self.outputlabel = Label(self.init_window_name, text="结果输出：", font=("宋体", 12, "bold"))
        self.outputlabel.grid(row=0, column=6, sticky='e')        
        self.Aslabel = Label(self.init_window_name, text="板受力筋计算面积As(mm\u00b2)：")
        self.Aslabel.grid(row=1, column=7, sticky='e')    
        self.AsText = Entry(self.init_window_name, width=20,textvariable="") 
        self.AsText.grid(row=1, column=8)
        self.AsText.insert(0,'250')
        self.Aspjlabel = Label(self.init_window_name, text="板受力筋配筋（eg:E8@200)：")
        self.Aspjlabel.grid(row=2, column=7, sticky='e')    
        self.AspjText = Entry(self.init_window_name, width=20,textvariable="") 
        self.AspjText.grid(row=2, column=8)
        self.Asmjlabel = Label(self.init_window_name, text="实际面积As(mm\u00b2)：")
        self.Asmjlabel.grid(row=3, column=7, sticky='e')    
        self.AsmjText = Entry(self.init_window_name, width=20,textvariable="") 
        self.AsmjText.grid(row=3, column=8)
        
        self.LAsvslabel = Label(self.init_window_name, text="梁箍筋计算面积（mm\u00b2/100mm）：")
        self.LAsvslabel.grid(row=5, column=7, sticky='e')    
        self.LAsvsText = Entry(self.init_window_name, width=20,textvariable="") 
        self.LAsvsText.grid(row=5, column=8)
        self.LAsvsText.insert(0,'250')
        self.LAsvspjlabel = Label(self.init_window_name, text="梁箍筋实际配筋（eg:E8@200）：")
        self.LAsvspjlabel.grid(row=6, column=7, sticky='e')    
        self.LAsvspjText = Entry(self.init_window_name, width=20,textvariable="") 
        self.LAsvspjText.grid(row=6, column=8)
        self.Lasvsmjlabel = Label(self.init_window_name, text="梁箍筋实际面积Asv/s(mm\u00b2/100mm)：")
        self.Lasvsmjlabel.grid(row=7, column=7, sticky='e')    
        self.LasvsmjText = Entry(self.init_window_name, width=20,textvariable="") 
        self.LasvsmjText.grid(row=7, column=8)
        
        self.AsLlabel = Label(self.init_window_name, text="梁上下纵筋计算面积As=As\'(mm\u00b2)：")
        self.AsLlabel.grid(row=9, column=7, sticky='e')    
        self.AsLText = Entry(self.init_window_name, width=20,textvariable="") 
        self.AsLText.grid(row=9, column=8)
        self.AsLText.insert(0,'400')
        self.LAspjlabel = Label(self.init_window_name, text="梁上下纵筋配筋（eg:2E16)：")
        self.LAspjlabel.grid(row=10, column=7, sticky='e')    
        self.LAspjText = Entry(self.init_window_name, width=20,textvariable="") 
        self.LAspjText.grid(row=10, column=8)
        self.LAsmjlabel = Label(self.init_window_name, text="梁上下纵筋实际面积As(mm\u00b2)：")
        self.LAsmjlabel.grid(row=11, column=7, sticky='e')    
        self.LAsmjText = Entry(self.init_window_name, width=20,textvariable="") 
        self.LAsmjText.grid(row=11, column=8)
        
        self.Astllabel = Label(self.init_window_name, text="梁抗扭纵筋计算面积Astl(mm\u00b2)：")
        self.Astllabel.grid(row=13, column=7, sticky='e')    
        self.AstlText = Entry(self.init_window_name, width=20,textvariable="") 
        self.AstlText.grid(row=13, column=8)
        self.AstlText.insert(0,'400')
        self.LAstlpjlabel = Label(self.init_window_name, text="梁抗扭纵筋配筋（eg:2E16)：")
        self.LAstlpjlabel.grid(row=14, column=7, sticky='e')    
        self.LAstlpjText = Entry(self.init_window_name, width=20,textvariable="") 
        self.LAstlpjText.grid(row=14, column=8)
        self.LAstlmjlabel = Label(self.init_window_name, text="梁抗扭纵筋实际面积As(mm\u00b2)：")
        self.LAstlmjlabel.grid(row=15, column=7, sticky='e')    
        self.LAstlmjText = Entry(self.init_window_name, width=20,textvariable="") 
        self.LAstlmjText.grid(row=15, column=8)   

        #按钮
        self.cal = Button(self.init_window_name, text="计算", bg="lightblue", width=10,command=lambda:self.calculation(self.hText.get(),self.lText.get(),self.bText.get(),self.dlText.get(),self.llText.get(),self.blText.get(),self.fcText.get(),self.ftText.get(),self.fyText.get(),self.pminText.get(),self.asText.get(),self.y0Text.get(),self.bbselection.get(),self.LbText.get(),self.LhText.get(),self.LLText.get()))  # 调用内部方法  加()为直接调用
        self.cal.grid(row=21, column=2)  
        self.docx = Button(self.init_window_name, text="生成计算书", bg="lightblue", width=10,command=lambda:self.txt(self.nameText.get(),self.hText.get(),self.lText.get(),self.bText.get(),self.dlText.get(),self.llText.get(),self.blText.get(),self.cgText.get(),self.fcText.get(),self.ftText.get(),self.sgText.get(),self.fyText.get(),self.pminText.get(),self.asText.get(),self.y0Text.get(),self.bbselection.get(),self.LbText.get(),self.LhText.get(),self.LLText.get()))  # 调用内部方法  加()为直接调用
        self.docx.grid(row=21, column=8)

    def calculation(self,h,l,b,dl,ll,bl,fc,ft,fy,pmin,a_s,y0,bool_beam,Lb,Lh,LL):
        h = float(h)  # 板厚 mm
        l = float(l)  # 挑出长度  mm
        b = float(b)  # 宽度  mm
        dl = float(dl)  # 恒载 kN/m\u00b2
        ll = float(ll)  # 活载  kN/m\u00b2
        bl = float(bl)  # 边缘翻边线荷载  kN/m\u00b2
        fc = float(fc)  # N/mm\u00b2
        ft = float(ft)  # N/mm\u00b2
        fy = float(fy)  # N/mm\u00b2
        pmin = float(pmin)  # 最小配筋率
        a_s = float(a_s)  # 纵筋合力点至近边距离
        y0 = float(y0)  # 重要性系数
        Lb = float(Lb)  # 梁宽mm  梁即使不算，也保留勿删
        Lh = float(Lh)  # 梁高mm
        LL = float(LL)  # 梁长mm

        # 板荷载计算
        Md = 0.5 * (25 * h / 1000 + dl) * l * l / 1000000
        Vd = (h / 1000 * 25 + dl) * l / 1000
        Ml = 0.5 * (25 * h / 1000 + ll) * l * l / 1000000
        Vl = (h / 1000 * 25 + ll) * l / 1000
        Mb = bl * l / 1000
        Vb = bl
        M = 1.3 * (Md + Mb) + 1.5 * Ml
        V = 1.3 * (Vd + Vb) + 1.5 * Vl
        # 内梁荷载计算
        Lx = V
        LTm = M + Lx * Lb / 2 / 1000
        Tmax = LTm * LL / 2 / 1000
        Mmax = 1 / 8 * Lx * LL * LL / 1000000
        Vmax = 0.5 * Lx * LL / 1000
        # 正截面受弯计算
        h0 = h - a_s
        xb = 0.518
        as1 = y0 * M * 1000000 / (fc * 1000 * h0 * h0)
        x1 = 1 - math.sqrt(1 - 2 * as1)  # 计算相对受压区高度
        if x1 > xb:
            tkinter.messagebox.showinfo('警告','受压区高度' + str(x1) + '>' + str(xb) + ';不满足要求')
        else:
            pass
        As = fc * b * h0 * x1 / fy
        p = As / (b * h)
        Asmin = pmin * b * h
        if As <Asmin:
            As = Asmin
        else:
            pass
        self.AsText.delete(0,END)
        self.AsText.insert(0,int(As))

        # 梁计算参数
        ast = 35
        asb = 35
        Lh0 = Lh - asb
        Wt = Lb * Lb * (3 * Lh0 - Lb) / 6
        bcor = Lb - 50
        hcor = Lh - ast - asb
        Acor = bcor * hcor
        ucor = 2 * (bcor + hcor)

        if bool_beam == '是':  # 梁截面设计
            allow1 = 0.25 * 1 * fc
            allow2 = 0.2 * 1 * fc
            result1 = Vmax * 1000 / (Lb * Lh0) + Tmax * 1000000 / (0.8 * Lb * Lb * (3 * Lh - Lb) / 6)
            if Lh0 / Lb <= 4:  # 梁截面验算
                if result1 <= allow1:
                    result_jm = '梁截面满足要求'
                else:
                    tkinter.messagebox.showinfo('警告','梁截面不满足要求')
            elif 4 < Lh0 / Lb < 6:
                allow3 = allow1 * (6 - Lh0 / Lb) / 2 + allow2 * (Lh0 / Lb - 4)
                if result1 <= allow3:
                    result_jm = '梁截面满足要求'
                else:
                    tkinter.messagebox.showinfo('警告','梁截面不满足要求')
            else:
                tkinter.messagebox.showinfo('梁高宽比大于6，受扭构件的截面尺寸要求及扭曲截面承载力计算应符合专门规定。')
            result2 = Vmax * 1000 / (Lb * Lh0) + Tmax * 1000000 / (Lb * Lb * (3 * Lh - Lb) / 6)
            allow4 = 0.7 * ft

            if result2 < allow4: # 剪扭计算判定
                tkinter.messagebox.showinfo('满足要求，可不进行构件受减扭承载力计算，但需满足规范构造要求。')
            else:
                pass
# 正截面受弯计算
            x = 2 * ast  
            AsL = (Mmax - 1 * fc * Lb * x * (Lh0 - x / 2)) / (fy * (Lh0 - ast))
            AsLmin = 0.002 * Lb * Lh
            if AsL < AsLmin:
                AsL = AsLmin
            else:
                pass
            self.AsLText.delete(0,END)
            self.AsLText.insert(0,int(AsL))
# 剪扭计算
            bt = 1.5 / (1 + 0.5 * Vmax * Wt / (Tmax * 1000 * Lb * Lh0))
            if bt< 0.5:
                bt = 0.5
            elif bt > 1:
                bt = 1
            else:
                pass
            Ast1_st = (Tmax * 1000000 - 0.35 * bt * ft * Wt) / (1.2 * math.sqrt(1.2) * fy * Acor)
            nAsv1_sv = (Vmax - 0.7 * (1.5 - bt) * ft * Lb * Lh0) / (1.25 * fy * Lh0)
            if nAsv1_sv < 0:
                nAsv1_sv = 0
            else:
                pass
            Asv_s = nAsv1_sv / 2 + Ast1_st
            psvmin = 0.28 * ft / fy
            Asvmin_s = psvmin * Lb / 2
            if Asv_s < Asvmin_s:
                Asv_s = Asvmin_s
            else:
                pass

            Astl = 1.2 * fy * ucor / fy * (Ast1_st)
            T_Vb = Tmax * 1000 / Vmax / Lb
            if T_Vb > 2:
                T_Vb = 2
            else:
                pass
            ptlmin = 0.6 * math.sqrt(Tmax * 1000 / (Vmax * Lb)) * ft / fy
            Atlmin = ptlmin * Lb * Lh
            if Astl < Atlmin:
                Astl = Atlmin
            else:
                pass
            
            self.LAsvsText.delete(0,END)
            self.LAsvsText.insert(0,int(Asv_s * 100))
            self.AstlText.delete(0,END)
            self.AstlText.insert(0,int(Astl))
            
        else:
            pass

    def txt(self,name,h,l,b,dl,ll,bl,concrete_g,fc,ft,steel_g,fy,pmin,a_s,y0,bool_beam,Lb,Lh,LL):
        # 板参数定义
        h = int(h)  # 板厚 mm
        l = int(l)  # 挑出长度  mm
        b = int(b)  # 宽度  mm
        dl = float(dl)  # 恒载 kN/m\u00b2
        ll = float(ll)  # 活载  kN/m\u00b2
        bl = float(bl)  # 边缘翻边线荷载  kN/m\u00b2
        fc = float(fc)  # N/mm\u00b2
        ft = float(ft)  # N/mm\u00b2
        fy = float(fy)  # N/mm\u00b2
        pmin = float(pmin)  # 最小配筋率
        a_s = int(a_s)  # 纵筋合力点至近边距离
        y0 = float(y0)  # 重要性系数
        Lb = int(Lb)  # 梁宽mm  梁即使不算，也保留勿删
        Lh = int(Lh)  # 梁高mm
        LL = int(LL)  # 梁长mm

        # 板荷载计算
        Md = 0.5 * (25 * h / 1000 + dl) * l * l / 1000000
        Vd = (h / 1000 * 25 + dl) * l / 1000
        Ml = 0.5 * (25 * h / 1000 + ll) * l * l / 1000000
        Vl = (h / 1000 * 25 + ll) * l / 1000
        Mb = bl * l / 1000
        Vb = bl
        M = 1.3 * (Md + Mb) + 1.5 * Ml
        V = 1.3 * (Vd + Vb) + 1.5 * Vl
        # 内梁荷载计算
        Lx = V
        LTm = M + Lx * Lb / 2 / 1000
        Tmax = LTm * LL / 2 / 1000
        Mmax = 1 / 8 * Lx * LL * LL / 1000000
        Vmax = 0.5 * Lx * LL / 1000
        # 正截面受弯计算
        h0 = h - a_s
        xb = 0.518
        as1 = y0 * M * 1000000 / (fc * 1000 * h0 * h0)
        x1 = 1 - math.sqrt(1 - 2 * as1)  # 计算相对受压区高度
        As = fc * b * h0 * x1 / fy
        p = As / (b * h)
        Asmin = pmin * b * h
        # 梁计算参数
        ast = 35
        asb = 35
        Lh0 = Lh - asb
        Wt = Lb * Lb * (3 * Lh0 - Lb) / 6
        bcor = Lb - 50
        hcor = Lh - ast - asb
        Acor = bcor * hcor
        ucor = 2 * (bcor + hcor)

        if bool_beam == '是':  # 梁截面设计
            allow1 = 0.25 * 1 * fc
            allow2 = 0.2 * 1 * fc
            allow3 = allow1 * (6 - Lh0 / Lb) / 2 + allow2 * (Lh0 / Lb - 4)
            result1 = Vmax * 1000 / (Lb * Lh0) + Tmax * 1000000 / (0.8 * Lb * Lb * (3 * Lh - Lb) / 6)
            result2 = Vmax * 1000 / (Lb * Lh0) + Tmax * 1000000 / (Lb * Lb * (3 * Lh - Lb) / 6)
            allow4 = 0.7 * ft
            bt = 1.5 / (1 + 0.5 * Vmax * Wt / (Tmax * 1000 * Lb * Lh0))
            
            x = 2 * ast
            AsL = (Mmax - 1 * fc * Lb * x * (Lh0 - x / 2)) / (fy * (Lh0 - ast))
            AsLmin = 0.002 * Lb * Lh
        else:
            pass

        # 文本生成
        document = Document()
        document.styles['Normal'].font.name = 'Times New Roman'  # 默认西体
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 默认中体
        document.styles['Normal'].font.size = Pt(10.5)  # 正文全局大小为10.5 小五
        document.styles['Heading 1'].font.size = Pt(14)  # 一级标题字体大小14 四号
        document.styles['Heading 1'].font.bold = True  # 加粗
        document.styles['Heading 2'].font.size = Pt(12)
        document.styles['Heading 2'].font.bold = True

        run = document.add_heading('悬挑板计算书', level=1)  # 标题
        run.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_heading('一、构件编号：' + name, level=2)  # 一、计算信息
        document.add_heading('二、计算信息', level=2)  # 计算信息
        document.add_paragraph('1. 几何参数\n        截面类型：矩形\n        截面宽度：b=' + str(b) + 'mm\n        截面高度：h=' + str(h) + 'mm')
        document.add_paragraph('2. 材料信息\n        混凝土等级：' + concrete_g + '  fc=' + str(fc) + 'N/mm\u00b2    ft=' + str(ft) + 'N/mm\u00b2\n        钢筋种类：' + steel_g + '   fy=' + str(fy) + 'N/mm\u00b2\n        最小配筋率：ρmin=' + str(pmin * 100) + '%\n        纵筋合力点至近边距离：as=' + str(a_s) + 'mm')
        document.add_paragraph('3. 设计参数\n        结构重要性系数：γo=' + str(y0))

        document.add_heading('三、荷载计算', level=2)  # 荷载计算
        document.add_paragraph('1. 恒载\n        板面及板底恒载DL=' + str(dl) + 'kN/m\u00b2\n        板端弯矩标准值Md=1/2*(25*h/1000+dl)*l\u00b2/10^6=' + str('%.1f' % Md) + 'kN.m\n        板端剪力标准值Vd=(h/1000*25+dl)*l/1000=' + str('%.1f' % Vd) + 'kN')
        document.add_paragraph('2. 活载\n        板面活载LL=' + str(ll) + 'kN/m\u00b2\n        板端弯矩标准值Ml=1/2*(25*h/1000+ll)*l\u00b2/10^6=' + str('%.1f' % Ml) + 'kN.m\n        板端剪力标准值Vl=(h/1000*25+ll)*l/1000=' + str('%.1f' % Vl) + 'kN')
        document.add_paragraph('3. 翻边\n        均布线荷载BL=' + str(bl) + 'kN/m\u00b2\n        板端弯矩标准值Mb=bl*l/1000=' + str('%.1f' % Mb) + 'kN.m\n        板端剪力标准值Vb=bl=' + str('%.1f' % Vb) + 'kN')
        document.add_paragraph('4. 荷载组合\n        弯矩设计值M=1.3*(Md+Mb)+1.5*Ml=' + str('%.1f' % M) + 'kN.m\n        剪力设计值V=1.3*(Vd+Vb)+1.5*Vl=' + str('%.1f' % V) + 'kN')
        if bool_beam == 'y':
            document.add_paragraph('5. 内梁\n        折算每米均布线荷载Lx=V=' + str(Lx) + 'kN/m\n        折算每米扭矩荷载LT/m=M+Lx*Lb/2/1000=' + str('%.1f' % LTm) + 'kN.m/m\n        梁中最大扭矩Tmax=LT/m*LL/2/1000=' + str('%.1f' % Tmax) + 'kN.m\n        梁中最大弯矩Mmax=1/8*Lx*LL\u00b2/10^6=' + str('%.1f' % Mmax) + 'kN.m\n        梁端最大剪力Vmax=1/2*Lx*LL/1000=' + str('%.1f' % Vmax) + 'kN')
        else:
            pass

        document.add_heading('四、板正截面受弯计算', level=2)  # 板正截面受弯计算
        document.add_paragraph('1. 计算截面有效高度:\n        ho=h-as=' + str(h) + '-' + str(a_s) + '=' + str(h0) + 'mm')
        document.add_paragraph('2. 计算相对受压区高度:\n        ξb=β1/(1+fy/(Es×εcu))=0.8/(1+' + str(fy) + '/(2×10^5×0.0033))=' + str(xb))
        document.add_paragraph('3. 确定计算系数:\n        αs=γo*M/(α1*fc*b*ho*ho)=' + str('%.1f' % y0) + '*' + str('%.1f' % M) + '*10^6/(1.0*' + str(fc) + '*1000*' + str(b) + '*' + str(h0) + '*' + str(h0) + ')=' + str('%.1f' % as1))
        document.add_paragraph('4. 计算相对受压区高度:\n        ξ=1-sqrt(1-2αs)=1-sqrt(1-2*' + str('%.1f' % as1) + ')=' + str('%.3f' % x1) + '≤ξb=0.518    满足要求。')
        document.add_paragraph('5. 计算纵向受拉钢筋面积:\n        As＝α1*fc*b*ho*ξ/fy=1.0*' + str(fc) + '*' + str(b) + '*' + str(h0) + '*' + str('%.3f' % x1) + '/' + str(fy) + '=' + str('%.1f' % As) + 'mm\u00b2')
        document.add_paragraph('6. 验算最小配筋率:\n        ρ=As/(b*h)=' + str('%.1f' % As) + '/(' + str(b) + '*' + str(h) + ')=' + str('%.2f' % (p * 100)) + '%')
        if p < pmin:
            document.add_paragraph('       ρ=' + str('%.2f' % (p * 100)) + '%<ρmin=' + str('%.2f' % (pmin * 100)) + '%，不满足最小配筋率要求，\n        取As=ρmin*b*h=' + str('%.2f' % (pmin * 100)) + '%*' + str(b) + '*' + str(h) + '=' + str('%.2f' % Asmin) + 'mm\u00b2')
            As = Asmin
        else:
            document.add_paragraph('       ρ=' + str('%.2f' % (p * 100)) + '%>ρmin=' + str('%.2f' % (pmin * 100)) + '%，满足最小配筋率要求。')

        Aspj = self.AspjText.get()
        Asmj = self.AsmjText.get()
        document.add_heading('五、板配筋', level=2)  # 板配筋
        document.add_paragraph('       实配为' + Aspj + '，实际配筋面积为' + Asmj + 'mm\u00b2。')

        if bool_beam == '是':
            document.add_heading('六：内梁计算', level=2)  # 内梁弯剪扭计算
# 1. 已知条件
            document.add_paragraph('1. 已知条件及计算要求：\n（1）已知条件：\n        计算截面：矩形梁    Lb=' + str(Lb) + 'mm,Lh=' + str(Lh) + 'mm。\n        砼等级：' + concrete_g + '，fc=' + str(fc) + 'N/mm\u00b2，ft=' + str(ft) + 'N/mm\u00b2\n        纵筋等级：' + steel_g + '，fy=' + str(fy) + 'N/mm\u00b2，fy\'=' + str(fy) + 'N/mm\u00b2\n        箍筋等级：' + steel_g + '，fy=' + str(fy) + 'N/mm\u00b2。')
            document.add_paragraph('      弯矩设计值M=' + str('%.1f' % Mmax) + 'kN.m，剪力设计值V=' + str('%.1f' % Vmax) + 'kN，扭矩设计值T=' + str('%.1f' % Tmax) + 'kN.m。')
            document.add_paragraph('（2）计算要求：\n        1）正截面受弯承载力计算\n        2）斜截面受剪承载力计算\n        3）受扭承载力计算')
# 2. 截面特征量计算
            document.add_paragraph('2. 截面特征量计算：\n        Wt=Lb*Lb*(3*Lh0-Lb)/6=' + str(Lb) + '*' + str(Lb) + '*(3*' + str(Lh0) + '-' + str(Lb) + ')/6=' + str('%.1f' % Wt) + 'mm^3\n        bcor=Lb-50=' + str(Lb) + '-50=' + str(bcor) + 'mm\n        hcor=Lh-as上-as下=' + str(Lh) + '-' + str(ast) + '-' + str(asb) + '=' + str(hcor) + 'mm\n        Acor=bcor*hcor=' + str(bcor) + '*' + str(hcor) + '=' + str(Acor) + 'mm\u00b2\n        ucor=2(bcor+hcor)=2*(' + str(bcor) + '+' + str(hcor) + ')=' + str(ucor) + 'mm')
# 3. 截面尺寸复核
            document.add_paragraph('3. 截面尺寸复核：')
            if Lh0 / Lb <= 4:
                document.add_paragraph('        Lh0/Lb=' + str('%.1f' % (Lh0 / Lb)) + '<4\n        Vmax/(Lb*h0)+Tmax/Wt=' + str('%.2f' % result1) + 'N/mm\u00b2<\n                0.25*βc*fc=' + str(0.25 * 1 * fc) + 'N/mm\u00b2\n        截面尺寸复核要求。')
            else:
                document.add_paragraph('        4<Lh0/Lb=' + str('%.1f' % (Lh0 / Lb)) + '<6\n        Vmax/(Lb*h0)+Tmax/(0.8*Lb*Lb*(3*Lh-Lb)/6)=' + str('%.2f' % result1) + 'N/mm\u00b2\n                <(6-Lh0/Lb)/2*0.25*1*fc)+(Lh0/Lb-4)*(0.2*1*fc)=' + str('%.2f' % allow3) + 'N/mm\u00b2\n        截面尺寸符合要求。')
# 4. 是否按剪扭计算
            document.add_paragraph('4. 是否需按剪扭计算:')
            if result2 < allow4:
                document.add_paragraph('        Vmax/(Lb*Lh0)+Tmax/Wt=' + str('%.1f' % Vmax) + '*1000/(' + str(Lb) + '*' + str(Lh0) + ')+' + str('%.1f' % Tmax) + '*1000000/' + str('%.1f' % Wt) + '=' + str('%.1f' % result2) + '\n                <0.7ft=0.7*' + str(ft) + '=' + str('%.1f' % allow4) + 'N/mm\u00b2\n        满足要求，可不进行构件受减扭承载力计算，但需满足规范构造要求。')
                tkinter.messagebox.showinfo('本计算小程序仅针对需要按剪扭计算进行计算，若不需要，可将第6.7节剪扭计算部分删除')
            else:
                document.add_paragraph('        Vmax/(Lb*Lh0)+Tmax/Wt=' + str('%.1f' % Vmax) + '*1000/(' + str(Lb) + '*' + str(Lh0) + ')+' + str('%.1f' % Tmax) + '*1000000/' + str('%.1f' % Wt) + '=' + str('%.1f' % result2) + '\n                >0.7ft=0.7*' + str(ft) + '=' + str('%.1f' % allow4) + 'N/mm\u00b2\n        需进行构件受减扭承载力计算。')
    # 5. 双筋梁正截面受弯计算
            document.add_paragraph('5. 正截面受弯承载力计算：\n（1）按双筋计算：as下=' + str(asb) + 'mm，as上=' + str(ast) + 'mm\n        假设采用上下对称配筋，则x=(fy*As-fy\'*As\')/(α1*fc*Lb)=0\n        则x≤ξb*Lh0，且需x≥2*a\'，所以取x=2*' + str(ast) + '=' + str(2 * ast) + 'mm\n        As=As\'=(M-α1*fc*Lb*x*(Lh0-x/2))/(fy\'*(Lh0-as\'))=(' + str('%.1f' % M) + '-1.0*' + str(fc) + '*' + str(Lb) + '*' + str(2 * ast) + '*(' + str(Lh0) + '-' + str(2 * ast) + '/2))/(' + str(fy) + '*(' + str(Lh0) + '-' + str(ast) + '))=' + str('%.1f' % AsL) +
                                'mm\u00b2\n（2）构造要求：非抗震ρmin=0.20%，Asmin=0.2%*' + str(Lb) + '*' + str(Lh) + '=' + str(AsLmin) + 'mm\u00b2')
            if AsL <= AsLmin:
                document.add_paragraph('（3）配筋面积：As = ' + str('%.1f' % AsL) + '≤Asmin=' + str(AsLmin) + 'mm\u00b2\n        需As=As\'=' + str(AsLmin) + 'mm\u00b2')
                AsL = AsLmin
            else:
                document.add_paragraph('（3）配筋面积：As = ' + str('%.1f' % AsL) + '>Asmin=' + str(AsLmin) + 'mm\u00b2\n        需As=As\'=' + str('%.1f' % AsL) + 'mm\u00b2')
    # 6. 确定箍筋
            if bt < 0.5:
                document.add_paragraph('6. 确定受剪及受扭箍筋\n（1）受扭所需箍筋量\n        计算剪扭构件受扭承载力降低系数βt\n        βt=1.5/(1+0.5*(V*Wt/(T*Lb*Lh0))=' + str('%.1f' % bt) + '<0.5\n        故取βt=0.5')
                bt = 0.5
            elif bt > 1:
                document.add_paragraph('6. 确定受剪及受扭箍筋\n（1）受扭所需箍筋量\n        计算剪扭构件受扭承载力降低系数βt\n        βt=1.5/(1+0.5*(V*Wt/(T*Lb*Lh0))=' + str('%.1f' % bt) + '>1.0\n        故取βt=1')
                bt = 1
            else:
                document.add_paragraph('6. 确定受剪及受扭箍筋\n（1）受扭所需箍筋量\n        计算剪扭构件受扭承载力降低系数βt\n        βt=1.5/(1+0.5*(V*Wt/(T*Lb*Lh0))=' + str('%.1f' % bt) + '>0.5且<1')
            Ast1_st = (Tmax * 1000000 - 0.35 * bt * ft * Wt) / (1.2 * math.sqrt(1.2) * fy * Acor)
            document.add_paragraph('        令配筋强度比ζ=1.2\n        Ast1/st=(T-0.35*βt*ft*Wt)/(1.2*sqrt(ζ)*fyv*Acor)=' + str('%.1f' % (Ast1_st * 100)) + 'mm\u00b2/100mm')
            nAsv1_sv = (Vmax - 0.7 * (1.5 - bt) * ft * Lb * Lh0) / (1.25 * fy * Lh0)
            Asv_s = nAsv1_sv / 2 + Ast1_st
            psvmin = 0.28 * ft / fy
            Asvmin_s = psvmin * Lb / 2
            if nAsv1_sv > 0:
                document.add_paragraph('（2）受剪所需箍筋数量\n        nAsv1/sv=(V-0.7(1.5-βt)*ft*Lb*Lh0)/(1.25*fyv*Lh0)=' + str('%.1f' % (nAsv1_sv * 100)) + 'mm\u00b2/100mm\n        按照双肢箍考虑，n=2,则单肢箍总的箍筋用量为\n        Asv1\"/s=Asv1/sv+Ast1/st=' + str('%.1f' % (Asv_s * 100)) + 'mm\u00b2/100mm\n        最小配箍率验算：\n        ρsv,min=0.28*ft/fy=' + str('%.2f' % (psvmin * 100)) + '%\n        单肢箍最小配箍面积：\n        Asv,min/s=ρsv,min*b/n=' + str('%.1f' % (Asvmin_s * 100)) + 'mm\u00b2/100mm')
            else:
                Asv_s = Ast1_st
                document.add_paragraph('（2）受剪所需箍筋数量\n        nAsv1/sv=(V-0.7(1.5-βt)*ft*Lb*Lh0)/(1.25*fyv*Lh0)=' + str('%.1f' % (nAsv1_sv * 100)) + 'mm\u00b2/100mm<0取0\n        按照双肢箍考虑，n=2,则单肢箍总的箍筋用量为\n        Asv1\"/s=Asv1/sv+Ast1/st=' + str('%.1f' % (Asv_s * 100)) + 'mm\u00b2/100mm\n        最小配箍率验算：\n        ρsv,min=0.28*ft/fy=' + str('%.2f' % (psvmin * 100)) + '%\n        单肢箍最小配箍面积：\n        Asv,min/s=ρsv,min*b/n=' + str('%.1f' % (Asvmin_s * 100)) + 'mm\u00b2/100mm')
            if Asv_s < Asvmin_s:
                document.add_paragraph('        因为Asv1\"/s<Asv,min/s,\n        故取Asv1\"/s=Asv,min/s=' + str('%.2f' % (psvmin * 100)) + 'mm\u00b2/100mm')
                Asv_s = Asvmin_s
            else:
                document.add_paragraph('        因此Asv1\"/s>Asv,min/s,满足最小配筋率要求。')
            Astl = 1.2 * fy * ucor / fy * (Ast1_st)
            # 7. 计算抗扭纵筋用量
            document.add_paragraph('7. 计算抗扭纵筋用量\n        Astl=ζ*fyv*ucor/fy*(Ast1/st)=' + str(int(Astl)) + 'mm\u00b2\n        抗扭纵筋最小配筋率：')
            T_Vb = Tmax * 1000 / Vmax / Lb
            if T_Vb < 2:
                document.add_paragraph('        T/(V*b)=' + str('%.1f' % T_Vb) + '<2')
            else:
                document.add_paragraph('        T/(V*b)=' + str('%.1f' % T_Vb) + '≥2，取T/(V*b)=2')
            ptlmin = 0.6 * math.sqrt(Tmax * 1000 / (Vmax * Lb)) * ft / fy
            Atlmin = ptlmin * Lb * Lh
            document.add_paragraph('        ρtl,min=0.6*sqrt(T/(V*b))*ft/fy=' + str('%.1f' % (ptlmin * 100)) + '%\n        抗扭纵筋最小配筋面积：Atl,min=ρtl,min*Lb*Lh=' + str(int(Atlmin)) + 'mm\u00b2')
            if Astl > Atlmin:
                document.add_paragraph('        因此Astl>Atl,min，满足最小配筋率要求。')
            else:
                document.add_paragraph('        因为Astl<Atl,min,故取Astl=Atl,min=' + str(int(Atlmin)))
                Astl = Atlmin
    # 8. 梁配筋
            document.add_paragraph('8.梁配筋')
            document.add_paragraph('        箍筋计算面积：Asv/s=' + str(int (Asv_s * 100)) + 'mm\u00b2/100mm')
            LAsvspj = self.LAsvspjText.get()
            Lasvsmj = self.LasvsmjText.get()
            document.add_paragraph('        选用：' + LAsvspj + '\n        实际Asv/s=' + Lasvsmj + 'mm\u00b2')
            document.add_paragraph('        顶部及底部纵筋计算面积：As=As\'=' + str(int( AsL)) + 'mm\u00b2')
            LAspj = self.LAspjText.get()
            Lasmj = self.LAsmjText.get()
            document.add_paragraph('        选用：' + LAspj + '\n        实际As=As\'=' + Lasmj + 'mm\u00b2')
            document.add_paragraph('        抗扭纵筋计算面积：Astl=' + str(int(AsL)) + 'mm\u00b2')
            LAstlpj = self.LAstlpjText.get()
            Lastlmj = self.LAstlmjText.get()
            document.add_paragraph('        选用：' + LAstlpj + '\n        实际Astl=' + Lastlmj + 'mm\u00b2')

        for paragraph in document.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)
            paragraph_format.line_spacing = Pt(18)

        document.save(name + '.docx')
        tkinter.messagebox.showinfo('Congratrulations!','计算书已成功导出！')
       
        
def gui_start():
    init_window = style.master              #实例化出一个父窗口
    ZMJ_PORTAL = MY_GUI(init_window)
    # 设置根窗口默认属性
    ZMJ_PORTAL.set_init_window()

    init_window.mainloop()          #父窗口进入事件循环，可以理解为保持窗口运行，否则界面不展示

gui_start()